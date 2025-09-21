#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
二维码处理核心模块
"""

import pandas as pd
import qrcode
from PIL import Image, ImageDraw, ImageFont
import os
import sys
import math
import concurrent.futures
import time
from typing import List, Tuple, Dict

# 尝试导入python-docx库
try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 添加src目录到Python路径
src_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(src_dir)

# 从core模块导入config
from core.config import *
from core.config import calculate_a4_layout

class QRCodeProcessor:
    """
    二维码处理核心类，提供二维码生成和A4图片合成的核心功能
    """
    
    def __init__(self):
        self.logger = self._get_logger()
        self.stop_event = None  # 用于取消操作的事件标志
        # 创建可重用的线程池，避免每次调用方法时重复创建
        self.qr_thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS)
        self.image_thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=MAX_IMAGE_WORKERS)
    
    def _get_logger(self):
        """获取日志记录器"""
        # 简单的日志记录实现，实际项目中可以使用更完善的日志系统
        return {
            'info': print,
            'error': print,
            'debug': print
        }
        
    def set_logger(self, logger_callback):
        """设置自定义日志记录回调函数"""
        self.logger = {
            'info': logger_callback,
            'error': logger_callback,
            'debug': logger_callback
        }
    
    def read_excel_in_batches(self, file_path: str, start_row: int, batch_size: int = BATCH_SIZE_EXCEL) -> List[str]:
        """
        分批读取Excel文件，避免内存溢出
        
        Args:
            file_path (str): Excel文件路径
            start_row (int): 开始读取的行数
            batch_size (int): 每批读取的行数
        
        Returns:
            List[str]: 读取到的字符串列表
        """
        all_strings = []
        # 计算需要跳过的行数（pandas从0开始计数）
        skip_rows = start_row - 1 if start_row > 1 else 0
        
        try:
            # 使用ExcelFile对象打开文件，并明确指定引擎为openpyxl
            xl = pd.ExcelFile(file_path, engine='openpyxl')
            sheet_name = xl.sheet_names[0]  # 使用第一个工作表
            
            # 获取工作表对象
            sheet = xl.parse(sheet_name)
            
            # 计算数据总行数
            total_rows = len(sheet)
            
            # 逐批读取数据
            for i in range(skip_rows, total_rows, batch_size):
                end_row = min(i + batch_size, total_rows)
                # 读取当前批次的数据
                chunk = sheet.iloc[i:end_row]
                
                # 处理数据
                for _, row in chunk.iterrows():
                    # 检查第一个单元格是否有值
                    if pd.notna(row.iloc[0]):
                        all_strings.append(str(row.iloc[0]))
            
        except Exception as e:
            error_msg = ERROR_MESSAGES["EXCEL_ERROR"].format(str(e))
            self.logger['error'](error_msg)
            raise Exception(error_msg)
        
        return all_strings
    
    def create_qr_code(self, data: str, output_path: str) -> None:
        """
        创建高清二维码
        
        Args:
            data (str): 二维码中包含的数据
            output_path (str): 输出文件路径
        """
        qr = qrcode.QRCode(
            version=QR_VERSION,
            error_correction=QR_ERROR_CORRECTION,
            box_size=QR_BOX_SIZE,
            border=QR_BORDER,
        )
        qr.add_data(data)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color=QR_FILL_COLOR, back_color=QR_BACK_COLOR)
        # 保存高清二维码，提高DPI值
        img.save(output_path, dpi=(IMAGE_DPI, IMAGE_DPI))
    
    def generate_qr_code_worker(self, data_group: Tuple[str, str, int, int]) -> Tuple[str, int, int, int]:
        """
        线程工作函数，用于并行生成二维码
        
        Args:
            data_group (Tuple): 包含数据、输出目录和索引范围的元组
        
        Returns:
            Tuple: 包含二维码文件路径、索引范围和线程ID的元组
        """
        import threading
        data, output_dir, start_idx, end_idx = data_group
        # 将生成的单张二维码命名加上Excel的行编号
        qr_file = os.path.join(output_dir, f"qr_row_{start_idx}_{end_idx}.png")
        self.create_qr_code(data, qr_file)
        # 返回线程ID
        thread_id = threading.get_ident()
        return (qr_file, start_idx, end_idx, thread_id)
    
    def generate_qr_codes(self, strings: List[str], output_dir: str, progress_callback=None) -> List[Tuple]:
        """
        批量生成二维码
        
        Args:
            strings (List[str]): 要编码的字符串列表
            output_dir (str): 输出目录路径
            progress_callback (callable, optional): 进度更新回调函数，接收已完成批次数量作为参数
        
        Returns:
            List[Tuple]: 包含二维码文件路径和索引范围的元组列表
        """
        qr_files = []
        
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 准备工作任务
        tasks = []
        for i in range(0, len(strings), QR_PER_IMAGE):
            end_i = min(i + QR_PER_IMAGE, len(strings))
            group = strings[i:end_i]
            data = ";".join(group)
            tasks.append((data, output_dir, i+1, end_i))
        
        # 记录总批次数
        total_batches = len(tasks)
        self.logger['info'](INFO_MESSAGES["START_QR_GENERATION"].format(total_batches))
        
        # 分批提交任务到线程池，避免一次性创建过多任务
        start_time = time.time()
        
        # 使用有序字典来保存结果，确保顺序正确
        result_dict = {}
        
        # 提交所有任务到可重用的线程池
        future_to_idx = {
            self.qr_thread_pool.submit(self.generate_qr_code_worker, task): i 
            for i, task in enumerate(tasks)
        }
        
        # 收集结果
        batch_start_time = {}
        
        for future in concurrent.futures.as_completed(future_to_idx):
            idx = future_to_idx[future]
            
            # 记录批次开始时间
            if idx not in batch_start_time:
                batch_start_time[idx] = time.time()
            
            # 检查是否需要取消
            if self.stop_event and self.stop_event.is_set():
                # 取消所有未完成的任务
                for f in future_to_idx:
                    if not f.done():
                        f.cancel()
                break
                
            try:
                result = future.result()
                result_dict[idx] = result
                
                # 记录批次完成时间和信息
                batch_end_time = time.time()
                batch_time = batch_end_time - batch_start_time[idx]
                
                # 提取线程ID（如果结果包含）
                thread_id = "N/A"
                if len(result) > 3:
                    thread_id = result[3]
                
                # 格式化批次信息，不包含线程ID以避免误解
                batch_info = INFO_MESSAGES["BATCH_COMPLETED"].format(
                    idx + 1, 
                    len(strings[idx*QR_PER_IMAGE:min((idx+1)*QR_PER_IMAGE, len(strings))]), 
                    batch_time
                )
                self.logger['info'](batch_info)
                
                # 调用进度回调函数（如果提供）
                if progress_callback:
                    # 计算已完成的批次数
                    completed_batches = len([r for r in result_dict.values() if r is not None])
                    progress_callback(completed_batches)
                
            except concurrent.futures.CancelledError:
                self.logger['info'](f"任务 {idx} 已取消")
            except Exception as e:
                error_msg = ERROR_MESSAGES["QR_GENERATION_ERROR"].format(idx, str(e))
                self.logger['error'](error_msg)
        
        # 按原始顺序重建结果列表
        qr_files = [result_dict[i] for i in sorted(result_dict.keys())]
        
        end_time = time.time()
        info_msg = INFO_MESSAGES["QR_GENERATION_COMPLETE"].format(len(qr_files), end_time - start_time)
        self.logger['info'](info_msg)
        
        return qr_files
    
    def process_a4_page_worker(self, page_data: Tuple[List[Tuple[str, int, int]], str, int, int, int, int, str]) -> str:
        """
        线程工作函数，用于并行处理A4页面
        
        Args:
            page_data (Tuple): 包含二维码文件组、输出目录、索引和行列数的元组
        
        Returns:
            str: 生成的A4图片文件路径
        """
        qr_files_group, output_dir, start_i, end_i, rows, cols, title = page_data
        
        # 创建A4大小的白色背景图片
        a4_image = Image.new('RGB', (A4_WIDTH, A4_HEIGHT), color=BACKGROUND_COLOR)
        draw = ImageDraw.Draw(a4_image)
        
        # 添加标题（如果有）
        if title:
            try:
                # 根据600 DPI设置字体大小，使打印时字体高度为0.92cm
                # 计算公式：像素值 = 厘米值 / 2.54厘米/英寸 * DPI值
                # 0.92 cm / 2.54 cm/inch * 600 DPI ≈ 217 像素
                font_size = 217  # 标题字体大小，确保打印时高度为0.92cm
                # 尝试多种中文字体，确保在不同系统上都能正常显示中文
                for font_name in ['simhei.ttf', 'simkai.ttf', 'msyh.ttc', 'microsoftyahei.ttf', 'simsun.ttc']:
                    try:
                        font = ImageFont.truetype(font_name, font_size)
                        break
                    except:
                        continue
                else:
                    # 如果所有中文字体都尝试失败，回退到默认字体
                    font = ImageFont.load_default()
            except:
                # 如果出现其他异常，使用默认字体
                font = ImageFont.load_default()
            
            # 计算标题位置（居中）
            title_width, title_height = draw.textbbox((0, 0), title, font=font)[2:4]
            title_x = (A4_WIDTH - title_width) // 2
            title_y = MARGIN_PIXELS + 100  # 标题上方留出100像素的额外空白
            
            # 绘制标题
            draw.text((title_x, title_y), title, fill=TEXT_COLOR, font=font)
            
            # 为标题增加额外的上边距
            title_margin = title_height + 250  # 标题下方留出150像素的额外空白，增加与二维码之间的间隙
        else:
            title_margin = 0  # 没有标题时不需要额外边距
        
        # 计算二维码的位置，考虑标题占用的空间和底部间距
        available_width = A4_WIDTH - 2 * MARGIN_PIXELS
        available_height = A4_HEIGHT - 2 * MARGIN_PIXELS - title_margin - MARGIN_PIXELS  # 额外减去底部间距
        
        qr_width = available_width // cols
        qr_height = available_height // rows
        
        # 放置二维码 - 调整元组解构以适应包含线程ID的4元素元组
        for idx, (qr_file, start_num, end_num, _) in enumerate(qr_files_group):
            try:
                # 打开二维码图片
                qr_img = Image.open(qr_file)
                # 调整二维码大小，使用LANCZOS算法保持高质量
                qr_img = qr_img.resize((qr_width, qr_height), Image.Resampling.LANCZOS)
                
                # 计算位置，考虑标题占用的空间
                col = idx % cols
                row = idx // cols
                x = MARGIN_PIXELS + col * qr_width
                y = MARGIN_PIXELS + title_margin + row * qr_height
                
                # 粘贴二维码到A4图片
                a4_image.paste(qr_img, (x, y))
                
                # 去掉A4纸上底部的编号范围，保留二维码图片
                pass
                
            except Exception as e:
                self.logger['error'](f"处理二维码 {qr_file} 时出错: {e}")
        
        # 保存A4图片
        if qr_files_group:
            start_num = qr_files_group[0][1]
            end_num = qr_files_group[-1][2]
            output_file = os.path.join(output_dir, f"{start_num}-{end_num}.png")
            a4_image.save(output_file, dpi=(IMAGE_DPI, IMAGE_DPI), quality=IMAGE_QUALITY)
            return output_file
        
        return ""
    
    def create_a4_image(self, qr_files: List[Tuple[str, int, int]], output_dir: str, qr_length_cm: float = DEFAULT_QR_LENGTH, title: str = "物料S/N清单") -> None:
        """
        使用多线程并行生成A4大小的图片
        
        Args:
            qr_files (List[Tuple]): 二维码文件路径和索引范围的元组列表
            output_dir (str): 输出目录路径
            qr_length_cm (float): 二维码边长，单位厘米，默认为配置文件中的DEFAULT_QR_LENGTH
        """
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 计算基础行列数
        base_rows, cols = calculate_a4_layout(qr_length_cm)
        
        # 如果有标题，需要考虑标题占用的高度，调整行数
        # 这里估算标题高度，并适当减少行数以确保二维码底部有足够间距
        if title:
            # 估算标题区域高度（包括边距和留白）
            estimated_title_height = 217 + 100 + 250  # 字体大小 + 上方留白 + 下方留白
            # 计算考虑标题后的可用高度
            effective_height_with_title = A4_HEIGHT - 2 * MARGIN_PIXELS - estimated_title_height - MARGIN_PIXELS  # 额外减去底部间距
            # 重新计算行数
            qr_length_px = int(qr_length_cm / 2.54 * IMAGE_DPI)
            rows = effective_height_with_title // qr_length_px
            rows = max(1, rows)  # 确保至少有1行
        else:
            rows = base_rows
            
        # 计算每页二维码数量
        qr_per_page = rows * cols
        
        # 准备工作任务
        tasks = []
        for i in range(0, len(qr_files), qr_per_page):
            end_i = min(i + qr_per_page, len(qr_files))
            group = qr_files[i:end_i]
            tasks.append((group, output_dir, i, end_i, rows, cols))
        
        # 使用多线程并行处理A4页面
        start_time = time.time()
        
        # 提交所有任务到可重用的线程池
        future_to_idx = {
            self.image_thread_pool.submit(self.process_a4_page_worker, task + (title,)): i 
            for i, task in enumerate(tasks)
        }
        
        # 收集结果 - 使用列表存储结果，确保按照提交顺序处理
        results = [None] * len(future_to_idx)
        for future in concurrent.futures.as_completed(future_to_idx):
            # 检查是否需要取消
            if self.stop_event and self.stop_event.is_set():
                # 取消所有未完成的任务
                for f in future_to_idx:
                    if not f.done():
                        f.cancel()
                break
                
            idx = future_to_idx[future]
            try:
                results[idx] = future.result()
            except concurrent.futures.CancelledError:
                self.logger['info'](f"A4图片任务 {idx} 已取消")
            except Exception as e:
                error_msg = ERROR_MESSAGES["IMAGE_GENERATION_ERROR"].format(idx, str(e))
                self.logger['error'](error_msg)
        
        # 按照提交顺序处理结果，确保二维码排列顺序与单线程一致
        for result in results:
            if result:
                success_msg = SUCCESS_MESSAGES["FILE_GENERATED"].format(result)
                self.logger['info'](success_msg)
                
    def create_docx_document(self, qr_files: List[Tuple[str, int, int]], output_dir: str, qr_length_cm: float = DEFAULT_QR_LENGTH, title: str = "物料S/N清单") -> str:
        """
        创建Word文档，将二维码以表格形式排列，方便用户自行排版
        
        Args:
            qr_files (List[Tuple]): 二维码文件路径和索引范围的元组列表
            output_dir (str): 输出目录路径
            qr_length_cm (float): 二维码边长，单位厘米
            title (str): 文档标题
            
        Returns:
            str: 生成的Word文档路径
        """
        if not DOCX_AVAILABLE:
            self.logger['error']("python-docx库未安装，无法生成Word文档")
            return ""
        
        try:
            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)
            
            # 创建新的Word文档
            doc = Document()
            
            # 设置页面边距为1厘米
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)
            
            # 添加标题
            if title:
                title_para = doc.add_heading(title, level=0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 计算每页二维码数量（以A4纸为基准）
            # 假设页面宽度约为21厘米，左右边距各1厘米，实际可用宽度约为19厘米
            page_width_cm = 19
            # 假设页面高度约为29.7厘米，上下边距各1厘米，实际可用高度约为27.7厘米
            page_height_cm = 27.7
            
            # 计算每行可以放置的二维码数量
            cols = int(page_width_cm // (qr_length_cm + 0.5))  # 0.5厘米间隔
            cols = max(1, cols)
            
            # 计算每列可以放置的二维码数量
            rows = int(page_height_cm // (qr_length_cm + 0.5))  # 0.5厘米间隔
            rows = max(1, rows)
            
            # 根据行列数计算每页二维码数量
            qr_per_page = cols * rows
            
            # 将二维码按页分组
            for page_idx in range(0, len(qr_files), qr_per_page):
                page_qr_files = qr_files[page_idx:page_idx + qr_per_page]
                
                # 添加分页符（除了第一页）
                if page_idx > 0:
                    doc.add_page_break()
                
                # 创建表格来放置二维码
                rows = math.ceil(len(page_qr_files) / cols)
                table = doc.add_table(rows=rows, cols=cols)
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 调整表格列宽
                for col in table.columns:
                    col.width = Cm(qr_length_cm + 0.5)
                
                # 填充表格
                for idx, qr_tuple in enumerate(page_qr_files):
                    # 解包元组，只获取前3个元素（忽略线程ID）
                    qr_file, start_idx, end_idx = qr_tuple[:3]
                    row_idx = idx // cols
                    col_idx = idx % cols
                    
                    # 获取单元格
                    cell = table.cell(row_idx, col_idx)
                    
                    # 在单元格中添加二维码图片
                    try:
                        # 计算图片在Word中的大小（厘米）
                        qr_length_inches = qr_length_cm / 2.54
                        cell.paragraphs[0].add_run().add_picture(qr_file, width=Inches(qr_length_inches))
                        
                        # 在图片下方添加编号（可选）
                        # run = cell.paragraphs[0].add_run(f"{start_idx}-{end_idx}")
                        # run.font.size = Pt(8)
                        
                        # 居中对齐
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        self.logger['error'](f"添加二维码 {qr_file} 到Word文档时出错: {e}")
                        continue
            
            # 保存Word文档
            output_file = os.path.join(output_dir, "二维码清单.docx")
            doc.save(output_file)
            
            self.logger['info'](f"Word文档已生成: {output_file}")
            return output_file
        except Exception as e:
            error_msg = f"生成Word文档时出错: {str(e)}"
            self.logger['error'](error_msg)
            return ""
    
    def shutdown(self):
        """
        关闭线程池，释放资源
        
        在线程池不再需要时调用此方法，确保资源被正确释放
        """
        # 关闭二维码生成线程池
        if hasattr(self, 'qr_thread_pool'):
            self.qr_thread_pool.shutdown(wait=True)
        
        # 关闭图像处理线程池
        if hasattr(self, 'image_thread_pool'):
            self.image_thread_pool.shutdown(wait=True)
            # 移除了错误的error_msg日志调用，因为error_msg只在异常情况下定义
        
        # 不记录完成时间，因为start_time变量在shutdown方法中未定义
        info_msg = INFO_MESSAGES["SHUTDOWN_COMPLETE"]
        self.logger['info'](info_msg)

# 创建全局实例，方便其他模块直接导入使用
qr_processor = QRCodeProcessor()

# 设置全局取消事件
def set_cancel_event(event):
    """设置取消事件标志"""
    qr_processor.stop_event = event

# 清除取消事件
def clear_cancel_event():
    """清除取消事件标志"""
    qr_processor.stop_event = None